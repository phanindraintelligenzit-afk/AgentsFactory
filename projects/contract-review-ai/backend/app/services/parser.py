import pdfplumber
import docx
import re
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from pathlib import Path


@dataclass
class ExtractedClause:
    clause_name: str
    clause_text: str
    start_pos: int
    end_pos: int
    confidence: float


class ClauseExtractor:
    """Extract clauses from contract text using regex patterns and heuristics"""
    
    # Standard clause patterns for NDA/MSA
    CLAUSE_PATTERNS = {
        "confidentiality": [
            r"(?i)(confidential|non.?disclosure|proprietary\s+information|trade\s+secrets?)",
        ],
        "term": [
            r"(?i)(term\s+of\s+agreement|duration|period\s+of\s+agreement|agreement\s+period)",
        ],
        "termination": [
            r"(?i)(terminat(?:e|ion)|expir(?:e|ation))",
        ],
        "liability_cap": [
            r"(?i)(limitation\s+of\s+liability|liability\s+cap|maximum\s+liability|cap\s+on\s+damages)",
        ],
        "indemnification": [
            r"(?i)(indemnif(?:y|ication)|hold\s+harmless)",
        ],
        "governing_law": [
            r"(?i)(governing\s+law|choice\s+of\s+law|jurisdiction)",
        ],
        "assignment": [
            r"(?i)(assign(?:ment)?|transfer|delegate)",
        ],
        "non_solicit": [
            r"(?i)(non.?solicit(?:ation)?)",
        ],
        "non_compete": [
            r"(?i)(non.?compete|compet(?:e|ition)\s+restrict)",
        ],
        "data_protection": [
            r"(?i)(data\s+protection|privacy|personal\s+data|gdpr|ccpa|data\s+processing)",
        ],
        "force_majeure": [
            r"(?i)(force\s+majeure|act\s+of\s+god)",
        ],
        "dispute_resolution": [
            r"(?i)(dispute\s+resolution|arbitration|mediation|litigation)",
        ],
        "ip_ownership": [
            r"(?i)(intellectual\s+property|ip\s+ownership|work\s+product|deliverables\s+ownership)",
        ],
        "payment_terms": [
            r"(?i)(payment\s+terms|fees|invoic(?:e|ing)|net\s+\d+)",
        ],
        "warranties": [
            r"(?i)(warrant(?:y|ies)|represent(?:ation|ations?))",
        ],
    }
    
    def __init__(self):
        # Compile patterns
        self.compiled_patterns = {}
        for clause_name, patterns in self.CLAUSE_PATTERNS.items():
            combined = "|".join(patterns)
            self.compiled_patterns[clause_name] = re.compile(combined, re.IGNORECASE)
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n\n".join(text_parts)
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        return "\n\n".join(text_parts)
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from file based on extension"""
        path = Path(file_path)
        if path.suffix.lower() == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif path.suffix.lower() in [".docx", ".doc"]:
            return self.extract_text_from_docx(file_path)
        else:
            # Try to read as text
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    
    def extract_clauses(self, text: str) -> List[Dict[str, Any]]:
        """Extract and identify clauses from contract text"""
        clauses = []
        
        # Split text into sections (by double newlines, numbered sections, etc.)
        sections = self._split_into_sections(text)
        
        for section in sections:
            if len(section.strip()) < 50:  # Skip very short sections
                continue
            
            # Identify clause type
            clause_name, confidence = self._identify_clause(section)
            
            if clause_name:
                clauses.append({
                    "clause_name": clause_name,
                    "clause_text": section.strip(),
                    "confidence": confidence,
                })
        
        # If no clauses found, treat entire text as one clause for each detected type
        if not clauses:
            detected = self._detect_clause_types(text)
            for clause_name, confidence in detected:
                clauses.append({
                    "clause_name": clause_name,
                    "clause_text": text[:3000],  # Limit length
                    "confidence": confidence,
                })
        
        return clauses
    
    def _split_into_sections(self, text: str) -> List[str]:
        """Split contract text into logical sections"""
        # Try multiple splitting strategies
        sections = []
        
        # Strategy 1: Double newlines
        double_newline_sections = [s.strip() for s in text.split("\n\n") if len(s.strip()) > 100]
        
        # Strategy 2: Numbered sections (1., 2., etc.)
        numbered_sections = re.split(r"\n\s*\d+\.\s+", text)
        numbered_sections = [s.strip() for s in numbered_sections if len(s.strip()) > 100]
        
        # Strategy 3: ALL CAPS headers
        caps_sections = re.split(r"\n\s*[A-Z][A-Z\s]{3,}:?\s*\n", text)
        caps_sections = [s.strip() for s in caps_sections if len(s.strip()) > 100]
        
        # Use the strategy that gives most reasonable sections
        all_sections = double_newline_sections + numbered_sections + caps_sections
        # Deduplicate
        seen = set()
        for s in all_sections:
            # Use first 100 chars as key
            key = s[:100]
            if key not in seen:
                seen.add(key)
                sections.append(s)
        
        return sections if sections else [text]
    
    def _identify_clause(self, text: str) -> tuple[str, float]:
        """Identify clause type from text"""
        scores = {}
        text_lower = text.lower()
        
        for clause_name, pattern in self.compiled_patterns.items():
            matches = pattern.findall(text)
            if matches:
                # Score based on number of matches and position (earlier = higher)
                first_match_pos = text_lower.find(matches[0].lower())
                position_score = 1.0 - (first_match_pos / len(text)) * 0.3
                scores[clause_name] = len(matches) * 0.3 + position_score
        
        if scores:
            best = max(scores, key=scores.get)
            return best, min(scores[best], 1.0)
        
        return "", 0.0
    
    def _detect_clause_types(self, text: str) -> List[tuple[str, float]]:
        """Detect all clause types present in text"""
        detected = []
        text_lower = text.lower()
        
        for clause_name, pattern in self.compiled_patterns.items():
            matches = pattern.findall(text)
            if matches:
                confidence = min(len(matches) * 0.2, 1.0)
                detected.append((clause_name, confidence))
        
        # Sort by confidence
        detected.sort(key=lambda x: x[1], reverse=True)
        return detected


class RedlineGenerator:
    """Generate redlined DOCX documents with tracked changes"""
    
    def __init__(self):
        pass
    
    def generate_redline(
        self,
        original_docx_path: str,
        analysis_results: List[Dict[str, Any]],
        output_path: str
    ) -> str:
        """Generate redlined DOCX with tracked changes"""
        doc = docx.Document(original_docx_path)
        
        # Process each analysis result
        for result in analysis_results:
            if result.get("risk_level") in ["high", "medium"] and result.get("redline_suggestion"):
                self._apply_redline(doc, result)
        
        doc.save(output_path)
        return output_path
    
    def _apply_redline(self, doc: docx.Document, result: Dict[str, Any]):
        """Apply redline to document for a specific clause"""
        clause_text = result.get("clause_text", "").strip()
        redline_suggestion = result.get("redline_suggestion", "")
        
        if not clause_text or not redline_suggestion:
            return
        
        # Find the clause text in document paragraphs
        for para in doc.paragraphs:
            if clause_text[:100] in para.text or para.text[:100] in clause_text:
                # Add comment with redline suggestion
                # Note: python-docx doesn't fully support tracked changes
                # We'll add a comment instead
                self._add_comment(para, f"REDLINE SUGGESTION ({result['risk_level'].upper()} RISK): {redline_suggestion}")
                break
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if clause_text[:100] in para.text or para.text[:100] in clause_text:
                            self._add_comment(para, f"REDLINE SUGGESTION ({result['risk_level'].upper()} RISK): {redline_suggestion}")
                            break
    
    def _add_comment(self, paragraph, comment_text: str):
        """Add comment to paragraph (limited support in python-docx)"""
        # python-docx has limited comment support
        # We'll add a run with distinctive formatting as a workaround
        run = paragraph.add_run(f"\n[AI REDLINE: {comment_text}]")
        run.font.color.rgb = docx.shared.RGBColor(0xFF, 0x00, 0x00)  # Red
        run.font.bold = True
        run.font.size = docx.shared.Pt(9)
    
    def generate_summary_docx(
        self,
        analysis_results: List[Dict[str, Any]],
        risk_summary: Dict[str, Any],
        contract_info: Dict[str, Any],
        output_path: str
    ) -> str:
        """Generate a summary DOCX with analysis results"""
        doc = docx.Document()
        
        # Title
        title = doc.add_heading("Contract Review Analysis Report", 0)
        title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Contract info
        doc.add_heading("Contract Information", level=1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        info = [
            ("Filename", contract_info.get("filename", "N/A")),
            ("Contract Type", contract_info.get("contract_type", "N/A")),
            ("Analysis Date", contract_info.get("analysis_date", "N/A")),
            ("Overall Risk Score", f"{risk_summary.get('overall_risk_score', 0)}/100"),
            ("Total Clauses Analyzed", str(risk_summary.get("total_clauses", 0))),
        ]
        for i, (key, value) in enumerate(info):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = str(value)
        
        # Risk Summary
        doc.add_heading("Risk Summary", level=1)
        risk_table = doc.add_table(rows=5, cols=2)
        risk_table.style = 'Table Grid'
        risks = [
            ("High Risk", risk_summary.get("high_risk", 0)),
            ("Medium Risk", risk_summary.get("medium_risk", 0)),
            ("Low Risk", risk_summary.get("low_risk", 0)),
            ("Approved", risk_summary.get("approved", 0)),
        ]
        for i, (key, value) in enumerate(risks):
            risk_table.rows[i].cells[0].text = key
            risk_table.rows[i].cells[1].text = str(value)
        # Overall score row
        risk_table.rows[4].cells[0].text = "Overall Risk Score"
        risk_table.rows[4].cells[1].text = f"{risk_summary.get('overall_risk_score', 0)}/100"
        
        # Detailed Analysis
        doc.add_heading("Detailed Clause Analysis", level=1)
        
        for result in analysis_results:
            risk_level = result.get("risk_level", "low").upper()
            clause_name = result.get("clause_name", "Unknown").replace("_", " ").title()
            
            # Color code by risk
            heading = doc.add_heading(f"{clause_name} — {risk_level}", level=2)
            
            if result.get("clause_text"):
                doc.add_paragraph("Original Clause:")
                p = doc.add_paragraph(result["clause_text"][:1000])
                p.style = 'Normal'
            
            if result.get("issues"):
                doc.add_paragraph("Issues Found:")
                for issue in result["issues"]:
                    doc.add_paragraph(issue, style='List Bullet')
            
            if result.get("redline_suggestion"):
                doc.add_paragraph("Suggested Redline:")
                p = doc.add_paragraph(result["redline_suggestion"])
                p.style = 'Normal'
                for run in p.runs:
                    run.font.color.rgb = docx.shared.RGBColor(0xFF, 0x00, 0x00)
            
            if result.get("explanation"):
                doc.add_paragraph("Analysis:")
                doc.add_paragraph(result["explanation"])
            
            doc.add_paragraph("")  # Spacer
        
        doc.save(output_path)
        return output_path


def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """Convert DOCX to PDF using LibreOffice (if available) or python-docx2pdf"""
    try:
        import subprocess
        result = subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf",
            "--outdir", str(Path(pdf_path).parent),
            docx_path
        ], capture_output=True, timeout=60)
        return result.returncode == 0
    except Exception:
        return False